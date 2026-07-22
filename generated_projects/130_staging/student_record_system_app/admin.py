from .models import Student, Course, Course, Enrollment, Enrollment
from django.contrib import admin
from django.http import HttpResponse
import csv
from docx import Document

class ExportAdminMixin:
    actions = ['export_as_csv', 'export_as_docx']

    def export_as_csv(self, request, queryset):
        meta = self.model._meta
        field_names = [field.name for field in meta.fields]
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename={meta.verbose_name_plural}.csv'
        writer = csv.writer(response)
        writer.writerow(field_names)
        for obj in queryset.select_related().iterator(chunk_size=1000):
            writer.writerow([getattr(obj, field) for field in field_names])
        return response
    export_as_csv.short_description = 'Export selected as CSV'

    def export_as_docx(self, request, queryset):
        meta = self.model._meta
        field_names = [field.name for field in meta.fields]
        document = Document()
        document.add_heading(str(meta.verbose_name_plural).title(), level=1)
        table = document.add_table(rows=1, cols=len(field_names))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i, name in enumerate(field_names):
            hdr_cells[i].text = name
        for obj in queryset.select_related().iterator(chunk_size=1000):
            row_cells = table.add_row().cells
            for i, name in enumerate(field_names):
                row_cells[i].text = str(getattr(obj, name))
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={meta.verbose_name_plural}.docx'
        document.save(response)
        return response
    export_as_docx.short_description = 'Export selected as Word document'

@admin.register(Student)
class StudentAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('first_name', 'last_name', 'email')
    list_filter = ('created_at',)
    search_fields = ('first_name', 'last_name', 'email')
    list_per_page = 25
    readonly_fields = ('created_at', 'updated_at')
    list_select_related = ()
    date_hierarchy = 'created_at'
    ordering = ('-id',)

@admin.register(Course)
class CourseAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('name', 'description')
    list_filter = ('created_at',)
    search_fields = ('name',)
    list_per_page = 25
    readonly_fields = ('created_at', 'updated_at')
    list_select_related = ()
    date_hierarchy = 'created_at'
    ordering = ('-id',)

@admin.register(Enrollment)
class EnrollmentAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('student', 'course', 'enrollment_date')
    list_filter = ('student', 'course')
    search_fields = ('student__first_name', 'student__last_name', 'course__name')
    list_per_page = 25
    readonly_fields = ('created_at', 'updated_at')
    list_select_related = ('student', 'course')
    date_hierarchy = 'enrollment_date'
    ordering = ('-id',)
@admin.register(Course)
class CourseAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('id', '__str__', 'created_at')
    search_fields = ('id',)
    list_per_page = 25

@admin.register(Enrollment)
class EnrollmentAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('id', '__str__', 'created_at')
    search_fields = ('id',)
    list_per_page = 25
