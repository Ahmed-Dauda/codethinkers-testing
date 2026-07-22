from django.contrib import admin
from django.http import HttpResponse
import csv
from docx import Document
from .models import Student, Attendance, Grade


class ExportAdminMixin:
    actions = ['export_as_csv', 'export_as_docx']

    def export_as_csv(self, request, queryset):
        meta = self.model._meta
        field_names = [field.name for field in meta.fields]
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename={meta.verbose_name_plural}.csv'
        writer = csv.writer(response)
        writer.writerow(field_names)
        # Use iterator() for memory-efficient export of large datasets
        for obj in queryset.select_related().iterator(chunk_size=1000):
            writer.writerow([getattr(obj, field) for field in field_names])
        return response
    export_as_csv.short_description = 'Export selected as CSV'

    def export_as_docx(self, request, queryset):
        meta = self.model._meta
        field_names = [field.name for field in meta.fields]
        document = Document()
        document.add_heading(str(meta.verbose_name_plural).title(), level=1)
        table = document.add_table(rows=1, cols=len(field_names))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i, name in enumerate(field_names):
            hdr_cells[i].text = name
        # Use iterator() for memory-efficient export
        for obj in queryset.select_related().iterator(chunk_size=1000):
            row_cells = table.add_row().cells
            for i, name in enumerate(field_names):
                row_cells[i].text = str(getattr(obj, name))
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename={meta.verbose_name_plural}.docx'
        document.save(response)
        return response
    export_as_docx.short_description = 'Export selected as Word document'


@admin.register(Student)
class StudentAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('name', 'email', 'created_at')
    list_filter = ('created_at',)
    search_fields = ('name', 'email')
    list_per_page = 25
    readonly_fields = ('created_at', 'updated_at') if hasattr(Student, 'updated_at') else ('created_at',)
    ordering = ('-id',)

    def get_queryset(self, request):
        return super().get_queryset(request).select_related().order_by('-id')


@admin.register(Attendance)
class AttendanceAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('student', 'date', 'status')
    list_filter = ('student', 'date', 'status')
    search_fields = ('student__name',)
    list_per_page = 25
    list_select_related = ('student',)
    date_hierarchy = 'date'
    autocomplete_fields = ('student',)
    ordering = ('-date',)

    def get_queryset(self, request):
        return super().get_queryset(request).select_related('student').order_by('-date')


@admin.register(Grade)
class GradeAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('student', 'subject', 'score')
    list_filter = ('student', 'subject')
    search_fields = ('student__name', 'subject')
    list_per_page = 25
    list_select_related = ('student',)
    autocomplete_fields = ('student',)
    ordering = ('-id',)

    def get_queryset(self, request):
        return super().get_queryset(request).select_related('student').order_by('-id')