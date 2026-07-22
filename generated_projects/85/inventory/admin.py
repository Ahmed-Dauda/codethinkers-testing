from django.contrib import admin
from django.http import HttpResponse
import csv
from docx import Document
from .models import Product, Category, Supplier, Stock, Sale

class ExportAdminMixin:
    actions = ['export_as_csv', 'export_as_docx']

    def export_as_csv(self, request, queryset):
        meta = self.model._meta
        field_names = [field.name for field in meta.fields]
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename={meta.verbose_name_plural}.csv'
        writer = csv.writer(response)
        writer.writerow(field_names)
        for obj in queryset:
            writer.writerow([getattr(obj, field) for field in field_names])
        return response
    export_as_csv.short_description = 'Export selected as CSV'

    def export_as_docx(self, request, queryset):
        meta = self.model._meta
        field_names = [field.name for field in meta.fields]
        document = Document()
        document.add_heading(str(meta.verbose_name_plural).title(), level=1)
        table = document.add_table(rows=1, cols=len(field_names))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i, name in enumerate(field_names):
            hdr_cells[i].text = name
        for obj in queryset:
            row_cells = table.add_row().cells
            for i, name in enumerate(field_names):
                row_cells[i].text = str(getattr(obj, name))
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={meta.verbose_name_plural}.docx'
        document.save(response)
        return response
    export_as_docx.short_description = 'Export selected as Word document'

@admin.register(Category)
class CategoryAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('name', 'created_at', 'updated_at')
    list_filter = ('created_at',)
    search_fields = ('name',)
    list_per_page = 25
    readonly_fields = ('created_at', 'updated_at')

@admin.register(Supplier)
class SupplierAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('name', 'created_at', 'updated_at')
    list_filter = ('created_at',)
    search_fields = ('name',)
    list_per_page = 25
    readonly_fields = ('created_at', 'updated_at')

@admin.register(Product)
class ProductAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('name', 'price', 'category', 'supplier', 'created_at')
    list_filter = ('category', 'supplier', 'created_at')
    search_fields = ('name',)
    list_per_page = 25
    readonly_fields = ('created_at', 'updated_at')
    list_select_related = ('category', 'supplier')
    autocomplete_fields = ('category', 'supplier')

@admin.register(Stock)
class StockAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('product', 'quantity', 'created_at')
    list_filter = ('product', 'created_at')
    search_fields = ('product__name',)
    list_per_page = 25
    readonly_fields = ('created_at', 'updated_at')
    list_select_related = ('product',)
    autocomplete_fields = ('product',)

@admin.register(Sale)
class SaleAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('product', 'quantity', 'date')
    list_filter = ('product', 'date')
    search_fields = ('product__name',)
    list_per_page = 25
    readonly_fields = ('date',)
    list_select_related = ('product',)
    autocomplete_fields = ('product',)