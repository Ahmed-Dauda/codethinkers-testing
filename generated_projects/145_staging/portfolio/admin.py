from django.contrib import admin
from django.http import HttpResponse
import csv
from docx import Document

class ExportAdminMixin:
    actions = ['export_as_csv', 'export_as_docx']

    def export_as_csv(self, request, queryset):
        meta = self.model._meta
        field_names = [field.name for field in meta.fields]
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename={meta.verbose_name_plural}.csv'
        writer = csv.writer(response)
        writer.writerow(field_names)
        for obj in queryset.select_related().iterator(chunk_size=1000):
            writer.writerow([getattr(obj, field) for field in field_names])
        return response
    export_as_csv.short_description = 'Export selected as CSV'

    def export_as_docx(self, request, queryset):
        meta = self.model._meta
        field_names = [field.name for field in meta.fields]
        document = Document()
        document.add_heading(str(meta.verbose_name_plural).title(), level=1)
        table = document.add_table(rows=1, cols=len(field_names))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i, name in enumerate(field_names):
            hdr_cells[i].text = name
        for obj in queryset.select_related().iterator(chunk_size=1000):
            row_cells = table.add_row().cells
            for i, name in enumerate(field_names):
                row_cells[i].text = str(getattr(obj, name))
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={meta.verbose_name_plural}.docx'
        document.save(response)
        return response
    export_as_docx.short_description = 'Export selected as Word document'

@admin.register(Skill)
class SkillAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('name', 'level')
    list_filter = ('level',)
    search_fields = ('name',)
    list_select_related = ()

@admin.register(Project)
class ProjectAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('title', 'url')
    list_filter = ()
    search_fields = ('title',)
    list_select_related = ()

@admin.register(Experience)
class ExperienceAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('company', 'role', 'start_date', 'end_date')
    list_filter = ('company',)
    search_fields = ('company', 'role')
    list_select_related = ()

@admin.register(ContactMessage)
class ContactMessageAdmin(ExportAdminMixin, admin.ModelAdmin):
    list_display = ('name', 'email', 'created_at')
    list_filter = ('created_at',)
    search_fields = ('name', 'email')
    list_select_related = ()